from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
import httpx
import json
from pydantic import BaseModel
from typing import Optional
import jwt
import bcrypt
from fastapi import Depends
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials

import os
import io
import base64
from datetime import datetime
import re
import pdfplumber
import docx
from pdf2image import convert_from_bytes
from PIL import Image
import cv2
import numpy as np
import chromadb

from pymongo import MongoClient
import gridfs
from bson import ObjectId
import requests
from dotenv import load_dotenv
from concurrent.futures import ThreadPoolExecutor, as_completed
import pytesseract

import platform

load_dotenv(override=True)

# OS Check for Deployment
if os.name == 'nt':
    # Local Windows Paths
    POPPLER_PATH = r"C:\Users\Dell\Desktop\intern\python_projects\RAGchatbot\poppler-26.02.0\Library\bin"
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
else:
    # Linux / Server Deployment Paths (detected automatically via system PATH)
    POPPLER_PATH = None

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

mongo_client = MongoClient(os.getenv("MONGO_URI"))
db = mongo_client["file_chatbot_db"]
fs = gridfs.GridFS(db)
files_collection = db["files_metadata"]
sessions_collection = db["chat_sessions"]
messages_collection = db["chat_messages"]


chroma_client = chromadb.PersistentClient(path="chroma_db")
collection = chroma_client.get_or_create_collection(name="documents")

GEMINI_API_KEY = os.getenv("GROQ_API_KEY") or os.getenv("GEMINI_API_KEY")
ADMIN_TOKEN = os.getenv("ADMIN_TOKEN", "supersecretadmin")

JWT_SECRET = os.getenv("JWT_SECRET", "super-secret-key-1234")
JWT_ALGORITHM = "HS256"

security = HTTPBearer()
users_collection = db["users"]

def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    token = credentials.credentials
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload.get("user_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token payload")
        return user_id
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Invalid token")

class UserCreate(BaseModel):
    username: str
    password: str
    
class UserLogin(BaseModel):
    username: str
    password: str

@app.post("/api/auth/signup")
def signup(user: UserCreate):
    existing = users_collection.find_one({"username": user.username})
    if existing:
        raise HTTPException(status_code=400, detail="Username already exists")
    
    hashed_pwd = bcrypt.hashpw(user.password.encode('utf-8'), bcrypt.gensalt())
    user_id = str(users_collection.insert_one({
        "username": user.username,
        "password": hashed_pwd,
        "created_at": datetime.utcnow()
    }).inserted_id)
    
    token = jwt.encode({"user_id": user_id}, JWT_SECRET, algorithm=JWT_ALGORITHM)
    return {"token": token, "username": user.username}

@app.post("/api/auth/login")
def login(user: UserLogin):
    db_user = users_collection.find_one({"username": user.username})
    if not db_user:
        raise HTTPException(status_code=400, detail="Invalid credentials")
        
    if not bcrypt.checkpw(user.password.encode('utf-8'), db_user["password"]):
        raise HTTPException(status_code=400, detail="Invalid credentials")
        
    token = jwt.encode({"user_id": str(db_user["_id"])}, JWT_SECRET, algorithm=JWT_ALGORITHM)
    return {"token": token, "username": user.username}



@app.on_event("startup")
def reindex_chromadb_on_startup():
    """Server restart hone par ChromaDB me wapas saara data daal do MongoDB se"""
    global collection
    ready_files = list(files_collection.find({"status": "ready"}))
    if not ready_files:
        print("✅ No files to re-index.")
        return

    reindexed = 0
    for file_doc in ready_files:
        file_id = str(file_doc["_id"])
        filename = file_doc.get("filename", "unknown")
        device_id = file_doc.get("device_id", "unknown")
        extracted_text = file_doc.get("extracted_text", "")

        if not extracted_text:
            continue

        # Check if chunks already exist in ChromaDB
        try:
            existing = collection.get(where={"file_id": file_id})
            if existing and existing["ids"] and len(existing["ids"]) > 0:
                continue  # Already indexed, skip
        except Exception:
            pass

        # Re-chunk and re-add (same helper jo upload ke waqt use hota hai)
        try:
            chunks_created = index_chunks_to_chroma(file_id, filename, device_id, extracted_text)
            if chunks_created:
                reindexed += 1
                print(f"🔄 Re-indexed: {filename} ({chunks_created} chunks)")
        except Exception as e:
            print(f"❌ Re-index failed for {filename}: {e}")

    print(f"✅ Startup re-index complete. {reindexed}/{len(ready_files)} files re-indexed.")


def table_to_text(table) -> str:
    text = ""
    for row in table:
        cleaned_row = [cell if cell else "" for cell in row]
        text += " | ".join(cleaned_row) + "\n"
    return text


def pil_image_to_base64(pil_image: Image.Image) -> str:
    buffered = io.BytesIO()
    # Ensure RGB for JPEG conversion
    pil_image.convert("RGB").save(buffered, format="JPEG")
    return base64.b64encode(buffered.getvalue()).decode("utf-8")


def run_ocr(pil_image):
    """Hybrid OCR: pehle Tesseract try karo (free, local, unlimited).
    Sirf tab Gemini Vision use karo jab Tesseract ka result bahut kam/khaali ho —
    isse Gemini ki chhoti free quota bachi rehti hai."""
    # Resize large images to max 1500px before OCR — speed boost
    max_dim = 1500
    if pil_image.width > max_dim or pil_image.height > max_dim:
        ratio = min(max_dim / pil_image.width, max_dim / pil_image.height)
        new_size = (int(pil_image.width * ratio), int(pil_image.height * ratio))
        pil_image = pil_image.resize(new_size, Image.LANCZOS)

    tesseract_text = ""
    try:
        gray = pil_image.convert("L")
        tesseract_text = pytesseract.image_to_string(gray).strip()
    except Exception as e:
        print("Tesseract Error:", e)

    # Agar Tesseract ne decent text nikal liya, wahi use karo — Gemini call bachao
    if len(tesseract_text) >= 15:
        return tesseract_text

    # Fallback: sirf tab Gemini Vision try karo jab Tesseract fail/khaali raha
    try:
        buffered = io.BytesIO()
        pil_image.save(buffered, format="PNG")
        img_str = base64.b64encode(buffered.getvalue()).decode("utf-8")

        gemini_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-3.5-flash:generateContent?key={GEMINI_API_KEY}"
        payload = {
            "contents": [{
                "parts": [
                    {"text": "Extract all text, code, and tables from this image exactly as written. Do not explain, just return the extracted text."},
                    {
                        "inlineData": {
                            "mimeType": "image/png",
                            "data": img_str
                        }
                    }
                ]
            }]
        }

        response = requests.post(gemini_url, json=payload, headers={"Content-Type": "application/json"}, timeout=30)

        if response.status_code == 429:
            # Gemini quota khatam — jo bhi Tesseract se mila usi pe guzara karo
            print("Gemini OCR quota exceeded, falling back to Tesseract result")
            return tesseract_text

        result = response.json()
        if "candidates" in result and result["candidates"]:
            text = result["candidates"][0].get("content", {}).get("parts", [{}])[0].get("text", "")
            return text.strip() or tesseract_text
        return tesseract_text
    except Exception as e:
        print("Gemini OCR Error:", e)
        return tesseract_text


def process_single_page(args):
    """Ek page process karta hai — parallel execution ke liye banaya gaya"""
    page_num, page_text, table_text, needs_ocr, ocr_image, img_boxes = args

    combined = (page_text + "\n" + table_text).strip()

    if needs_ocr and ocr_image is not None:
        if img_boxes and len(img_boxes) > 0:
            for box in img_boxes:
                scale_x = ocr_image.width / box['page_width']
                scale_y = ocr_image.height / box['page_height']
                crop_box = (
                    int(box['x0'] * scale_x),
                    int(box['top'] * scale_y),
                    int(box['x1'] * scale_x),
                    int(box['bottom'] * scale_y)
                )
                try:
                    cropped_img = ocr_image.crop(crop_box)
                    ocr_text = run_ocr(cropped_img)
                    if len(ocr_text.strip()) > 5:
                        combined = combined + "\n[IMAGE/CODE CONTENT]\n" + ocr_text.strip()
                except Exception as e:
                    print("Crop error:", e)
        else:
            ocr_text = run_ocr(ocr_image)
            if len(ocr_text.strip()) > 5:
                combined = combined + "\n[IMAGE/CODE CONTENT]\n" + ocr_text.strip()

    return page_num, combined


def extract_text_from_pdf_gemini(file_bytes: bytes, progress_callback=None) -> list:
    """PRIMARY: Gemini se seedha PDF ka text nikalo — 2M token context, native PDF support.
    Ek hi API call mein poora text aa jaata hai, bahut fast!"""
    try:
        pdf_base64 = base64.b64encode(file_bytes).decode("utf-8")
        
        if progress_callback:
            progress_callback(10)
        
        gemini_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-3.5-flash:generateContent?key={GEMINI_API_KEY}"
        payload = {
            "contents": [{
                "parts": [
                    {"text": "Extract ALL text content from this PDF document. Include all headings, paragraphs, tables, lists, captions, and any other textual content. Maintain the original structure and formatting as much as possible. For tables, format them with | separators. Do not summarize or skip any content — extract everything exactly as written. Return ONLY the extracted text, no explanations."},
                    {
                        "inlineData": {
                            "mimeType": "application/pdf",
                            "data": pdf_base64
                        }
                    }
                ]
            }],
            "generationConfig": {
                "maxOutputTokens": 65536
            }
        }

        if progress_callback:
            progress_callback(20)
        
        response = requests.post(gemini_url, json=payload, headers={"Content-Type": "application/json"}, timeout=120)
        
        if progress_callback:
            progress_callback(80)
        
        if response.status_code == 200:
            result = response.json()
            if "candidates" in result and result["candidates"]:
                text = result["candidates"][0].get("content", {}).get("parts", [{}])[0].get("text", "")
                if text and len(text.strip()) > 50:
                    print(f"✅ Gemini PDF extraction successful: {len(text)} chars")
                    if progress_callback:
                        progress_callback(100)
                    return [{"page": 1, "text": text.strip()}]
        
        print(f"⚠️ Gemini PDF extraction returned status {response.status_code}, falling back to pdfplumber")
        return None
        
    except Exception as e:
        print(f"⚠️ Gemini PDF extraction failed: {e}, falling back to pdfplumber")
        return None


def extract_text_from_pdf_pdfplumber(file_bytes: bytes, progress_callback=None) -> list:
    """FALLBACK: Pdfplumber + OCR pipeline — sirf tab use hota hai jab Gemini fail ho."""
    page_data = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        total_pages = len(pdf.pages)
        page_area_cache = {}

        # Step 1: First decide which pages need OCR (fast, sequential)
        needs_ocr_flags = []
        page_image_boxes = []
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            tables = page.extract_tables()
            table_text = ""
            for table in tables:
                table_text += "\n[TABLE]\n" + table_to_text(table) + "\n"

            combined = (page_text + "\n" + table_text).strip()
            
            boxes = []
            if len(page.images) > 0:
                for img in page.images:
                    if (img["width"] * img["height"]) > (page.width * page.height * 0.01):
                        boxes.append({
                            'x0': img['x0'], 'top': img['top'], 'x1': img['x1'], 'bottom': img['bottom'],
                            'page_width': float(page.width), 'page_height': float(page.height)
                        })

            has_no_text = len(combined) < 20

            # OCR sirf tab karo jab page mein text hi nahi hai
            if has_no_text:
                needs_ocr = True
            elif len(combined) < 50 and len(boxes) > 0:
                needs_ocr = True
            else:
                needs_ocr = False

            needs_ocr_flags.append(needs_ocr)
            page_image_boxes.append(boxes if needs_ocr else [])
            page_data.append((page_text, table_text))

        # Step 2: Convert ONLY the pages that need OCR — single batch call
        ocr_images_map = {}
        ocr_page_indices = [i for i, needs in enumerate(needs_ocr_flags) if needs]
        if ocr_page_indices:
            print(f"🔍 OCR needed for {len(ocr_page_indices)}/{total_pages} pages: {[p+1 for p in ocr_page_indices]}")
            try:
                first_p = min(ocr_page_indices) + 1
                last_p = max(ocr_page_indices) + 1
                if POPPLER_PATH:
                    all_range_images = convert_from_bytes(file_bytes, poppler_path=POPPLER_PATH, dpi=100,
                                                         first_page=first_p, last_page=last_p)
                else:
                    all_range_images = convert_from_bytes(file_bytes, dpi=100,
                                                         first_page=first_p, last_page=last_p)
                for page_idx in ocr_page_indices:
                    img_index = page_idx - (first_p - 1)
                    if 0 <= img_index < len(all_range_images):
                        ocr_images_map[page_idx] = all_range_images[img_index]
            except Exception as e:
                print(f"❌ Batch image conversion failed: {e}")
        else:
            print(f"✅ No OCR needed — all {total_pages} pages have text")

        # Step 3: Parallel OCR — 6 pages ek saath process honge
        results = {}
        completed = 0

        tasks = [
            (i, page_data[i][0], page_data[i][1], needs_ocr_flags[i], ocr_images_map.get(i), page_image_boxes[i])
            for i in range(total_pages)
        ]

        with ThreadPoolExecutor(max_workers=6) as executor:
            futures = {executor.submit(process_single_page, task): task[0] for task in tasks}

            for future in as_completed(futures):
                page_num, combined_text = future.result()
                results[page_num] = combined_text
                completed += 1

                if progress_callback:
                    percent = int((completed / total_pages) * 100)
                    progress_callback(percent)

        # Original page order me wapas jodo (keep page mapping)
        pages_data = [{"page": i+1, "text": results.get(i, "")} for i in range(total_pages)]
        return pages_data


def extract_text_from_pdf(file_bytes: bytes, progress_callback=None) -> list:
    """Smart extraction: Pehle Gemini try karo (fast, 2M tokens), fail ho toh pdfplumber fallback."""
    # Try Gemini first — single API call, bahut fast
    result = extract_text_from_pdf_gemini(file_bytes, progress_callback)
    if result:
        return result
    
    # Fallback to pdfplumber + OCR pipeline
    print("📄 Using pdfplumber fallback for PDF extraction...")
    return extract_text_from_pdf_pdfplumber(file_bytes, progress_callback)


def extract_images_from_docx(file_bytes: bytes):
    """DOCX ke andar embedded images (screenshots, diagrams, code snippets) nikalta hai"""
    images = []
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_bytes = rel.target_part.blob
                    img = Image.open(io.BytesIO(image_bytes))
                    images.append(img)
                except Exception as e:
                    print("DOCX image extract error:", e)
    except Exception as e:
        print("DOCX rels error:", e)
    return images


def process_docx_image(img):
    """Ek DOCX image ko OCR karta hai — chhoti images (icons/bullets) skip karta hai"""
    try:
        if img.width * img.height < 40000:  # ~200x200 se chhoti images ignore
            return ""
        ocr_text = run_ocr(img.convert("RGB"))
        return ocr_text.strip() if len(ocr_text.strip()) > 5 else ""
    except Exception as e:
        print("DOCX image OCR error:", e)
        return ""


def extract_text_from_docx(file_bytes: bytes, progress_callback=None) -> list:
    doc = docx.Document(io.BytesIO(file_bytes))
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        text += "\n[TABLE]\n"
        for row in table.rows:
            row_text = " | ".join([cell.text for cell in row.cells])
            text += row_text + "\n"

    # OCR fallback — DOCX ke andar embedded screenshots/diagrams bhi padho
    images = extract_images_from_docx(file_bytes)
    if images:
        completed = 0
        with ThreadPoolExecutor(max_workers=6) as executor:
            futures = [executor.submit(process_docx_image, img) for img in images]
            for future in as_completed(futures):
                ocr_text = future.result()
                if ocr_text:
                    text += "\n[IMAGE/CODE CONTENT]\n" + ocr_text
                completed += 1
                if progress_callback:
                    percent = int((completed / len(images)) * 100)
                    progress_callback(percent)

    return [{"page": 1, "text": text}]


def extract_text_from_image(file_bytes: bytes) -> list:
    img = Image.open(io.BytesIO(file_bytes))
    return [{"page": 1, "text": run_ocr(img)}]


def extract_text(file_bytes: bytes, filename: str, progress_callback=None) -> str:
    ext = filename.split(".")[-1].lower()

    if ext == "pdf":
        return extract_text_from_pdf(file_bytes, progress_callback)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes, progress_callback)
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")
    elif ext in ["jpg", "jpeg", "png"]:
        return extract_text_from_image(file_bytes)
    else:
        return ""


def chunk_text(pages_data, chunk_size: int = 500, overlap: int = 50):
    """Semantic chunking with page number tracking."""
    final_chunks = []
    
    # Backward compatibility for reindexing string payloads
    if isinstance(pages_data, str):
        pages_data = [{"page": 1, "text": pages_data}]

    for page_item in pages_data:
        text = page_item["text"]
        page_num = page_item["page"]
        
        paragraphs = re.split(r'\n\s*\n', text)
        chunks = []
        current_chunk = ""
        current_words = 0
        
        for p in paragraphs:
            p = p.strip()
            if not p:
                continue
                
            p_words = len(p.split())
            
            if current_words + p_words > chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                sentences = re.split(r'(?<=[.!?])\s+', current_chunk)
                overlap_text = ""
                overlap_words = 0
                for s in reversed(sentences):
                    s_words = len(s.split())
                    if overlap_words + s_words > overlap and overlap_text:
                        break
                    overlap_text = s + " " + overlap_text
                    overlap_words += s_words
                    
                current_chunk = overlap_text.strip()
                current_words = overlap_words
                
            current_chunk += "\n\n" + p if current_chunk else p
            current_words += p_words
            
        if current_chunk:
            chunks.append(current_chunk.strip())
            
        for c in chunks:
            words = c.split()
            if len(words) > chunk_size + 100:
                i = 0
                while i < len(words):
                    final_chunks.append({"text": " ".join(words[i:i + chunk_size]), "page": page_num})
                    i += chunk_size - overlap
            else:
                final_chunks.append({"text": c, "page": page_num})
                
    return final_chunks


@app.get("/")
def read_root():
    return {"message": "Backend is running 🚀"}


@app.get("/clear-all-data")
def clear_all_data(admin_token: str = Header(None)):
    if admin_token != ADMIN_TOKEN:
        raise HTTPException(status_code=403, detail="Forbidden: Invalid admin token")
    global collection
    # Clear MongoDB Collections
    db.drop_collection("files_metadata")
    db.drop_collection("chat_sessions")
    db.drop_collection("chat_messages")
    db.drop_collection("fs.files")
    db.drop_collection("fs.chunks")
    
    # Clear ChromaDB
    try:
        chroma_client.delete_collection(name="documents")
    except Exception:
        pass
    collection = chroma_client.get_or_create_collection(name="documents")
    
    return {"message": "All data (Files, Chats, ChromaDB) has been completely wiped. You can now start fresh."}


def compute_text_length(extracted) -> int:
    """extracted PDF/DOCX/image ke liye pages ki list hoti hai, TXT ke liye plain string.
    Dono cases me total character count nikalta hai."""
    if isinstance(extracted, str):
        return len(extracted)
    return sum(len(p.get("text", "")) for p in extracted)


def index_chunks_to_chroma(file_id: str, filename: str, device_id: str, extracted) -> int:
    """Chunking + embedding ka shared logic — upload aur startup reindex, dono isi se guzarte hain
    taaki dono jagah same (sahi) format use ho."""
    chunks_data = chunk_text(extracted)
    if not chunks_data:
        return 0

    chunks_text = [c["text"] for c in chunks_data]
    ids = [f"{file_id}_chunk_{i}" for i in range(len(chunks_text))]
    metadatas = [
        {
            "filename": filename,
            "file_id": file_id,
            "chunk_index": i,
            "device_id": device_id,
            "page_num": c["page"],
        }
        for i, c in enumerate(chunks_data)
    ]
    # ChromaDB ka default embedding function use ho raha hai (RAM bachane ke liye)
    collection.add(documents=chunks_text, ids=ids, metadatas=metadatas)
    return len(chunks_data)


def process_file_background(file_id: str, file_bytes: bytes, filename: str, device_id: str):
    def update_progress(percent):
        files_collection.update_one(
            {"_id": ObjectId(file_id)},
            {"$set": {"progress": percent}}
        )

    try:
        extracted = extract_text(file_bytes, filename, progress_callback=update_progress)
        chunks_created = index_chunks_to_chroma(file_id, filename, device_id, extracted)

        files_collection.update_one(
            {"_id": ObjectId(file_id)},
            {"$set": {
                "extracted_text": extracted,
                "status": "ready",
                "progress": 100,
                "text_length": compute_text_length(extracted),
                "chunks_created": chunks_created
            }}
        )
    except Exception as e:
        files_collection.update_one(
            {"_id": ObjectId(file_id)},
            {"$set": {"status": "error", "error_message": str(e)}}
        )


@app.post("/upload")
async def upload_file(file: UploadFile = File(...), background_tasks: BackgroundTasks = None, user_id: str = Depends(get_current_user)):
    file_bytes = await file.read()
    file_size = len(file_bytes)

    gridfs_id = fs.put(file_bytes, filename=file.filename, content_type=file.content_type)

    file_doc = {
        "filename": file.filename,
        "gridfs_id": gridfs_id,
        "file_size": file_size,
        "extracted_text": "",
        "status": "processing",
        "progress": 0,
        "upload_date": datetime.utcnow(),
        "device_id": user_id
    }
    result = files_collection.insert_one(file_doc)
    file_id = str(result.inserted_id)

    background_tasks.add_task(process_file_background, file_id, file_bytes, file.filename, user_id)

    return {
        "file_id": file_id,
        "filename": file.filename,
        "size": file_size,
        "message": "File uploaded successfully. Processing has started.",
        "status": "processing"
    }


@app.get("/files/{file_id}/status")
def get_file_status(file_id: str, user_id: str = Depends(get_current_user)):
    file_doc = files_collection.find_one({"_id": ObjectId(file_id), "device_id": user_id})
    if not file_doc:
        raise HTTPException(status_code=404, detail="File not found")
    return {
        "status": file_doc.get("status", "unknown"),
        "progress": file_doc.get("progress", 0),
        "chunks_created": file_doc.get("chunks_created", 0),
        "text_length": file_doc.get("text_length", 0),
        "error_message": file_doc.get("error_message", None)
    }


@app.get("/files")
def list_files(user_id: str = Depends(get_current_user)):
    files = files_collection.find({"device_id": user_id}, {"filename": 1, "file_size": 1, "upload_date": 1, "status": 1, "progress": 1})
    result = []
    for f in files:
        result.append({
            "file_id": str(f["_id"]),
            "filename": f["filename"],
            "size": f["file_size"],
            "upload_date": f["upload_date"].isoformat(),
            "status": f.get("status", "ready"),
            "progress": f.get("progress", 100)
        })
    return {"files": result}


@app.get("/files/{file_id}/download")
def download_file(file_id: str, user_id: str = Depends(get_current_user)):
    file_doc = files_collection.find_one({"_id": ObjectId(file_id), "device_id": user_id})
    if not file_doc:
        raise HTTPException(status_code=404, detail="File not found")
    grid_file = fs.get(file_doc["gridfs_id"])
    return StreamingResponse(
        io.BytesIO(grid_file.read()),
        media_type="application/octet-stream",
        headers={"Content-Disposition": f"attachment; filename={file_doc['filename']}"}
    )


@app.delete("/files/{file_id}")
def delete_file(file_id: str, user_id: str = Depends(get_current_user)):
    file_doc = files_collection.find_one({"_id": ObjectId(file_id), "device_id": user_id})
    if not file_doc:
        raise HTTPException(status_code=404, detail="File not found")
    fs.delete(file_doc["gridfs_id"])
    collection.delete(where={"file_id": file_id})
    files_collection.delete_one({"_id": ObjectId(file_id)})
    return {"message": "File deleted successfully"}


@app.post("/chats")
def create_chat_session(user_id: str = Depends(get_current_user)):
    session_doc = {"title": "New Chat", "created_at": datetime.utcnow(), "device_id": user_id}
    result = sessions_collection.insert_one(session_doc)
    return {"session_id": str(result.inserted_id), "title": "New Chat"}


@app.get("/chats")
def list_chat_sessions(user_id: str = Depends(get_current_user)):
    sessions = sessions_collection.find({"device_id": user_id}).sort("created_at", -1)
    result = []
    for s in sessions:
        result.append({
            "session_id": str(s["_id"]),
            "title": s["title"],
            "created_at": s["created_at"].isoformat()
        })
    return {"sessions": result}


@app.delete("/chats/{session_id}")
def delete_chat_session(session_id: str, user_id: str = Depends(get_current_user)):
    session_doc = sessions_collection.find_one({"_id": ObjectId(session_id), "device_id": user_id})
    if not session_doc:
        raise HTTPException(status_code=404, detail="Session not found")
    sessions_collection.delete_one({"_id": ObjectId(session_id)})
    messages_collection.delete_many({"session_id": session_id})
    return {"message": "Chat deleted successfully"}


@app.get("/chats/{session_id}/messages")
def get_session_messages(session_id: str, user_id: str = Depends(get_current_user)):
    session_doc = sessions_collection.find_one({"_id": ObjectId(session_id), "device_id": user_id})
    if not session_doc:
        raise HTTPException(status_code=404, detail="Session not found")
    messages = messages_collection.find({"session_id": session_id}).sort("timestamp", 1)
    result = []
    for msg in messages:
        result.append({"sender": msg["sender"], "text": msg["text"]})
    return {"messages": result}


class ChatRequest(BaseModel):
    session_id: str
    question: str
    file_id: Optional[str] = None  # Optional file_id for per-file filtering

@app.post("/chat")
async def chat(request: ChatRequest, user_id: str = Depends(get_current_user)):
    question = request.question
    session_id = request.session_id
    file_id = request.file_id

    session_doc = sessions_collection.find_one({"_id": ObjectId(session_id), "device_id": user_id})
    if not session_doc:
        raise HTTPException(status_code=404, detail="Session not found")

    # User message save karo
    messages_collection.insert_one({
        "session_id": session_id,
        "sender": "user",
        "text": question,
        "timestamp": datetime.utcnow()
    })

    # Title update (pehla message)
    existing_count = messages_collection.count_documents({"session_id": session_id})
    if existing_count == 1:
        title = question[:40] + ("..." if len(question) > 40 else "")
        sessions_collection.update_one(
            {"_id": ObjectId(session_id)},
            {"$set": {"title": title}}
        )

    # Multi-turn Query Rewriting
    past_messages = list(messages_collection.find({"session_id": session_id}).sort("timestamp", -1).limit(6))
    past_messages.reverse() # chronological order
    
    search_query = question
    if len(past_messages) > 1:
        # Ask Gemini to rewrite the query
        history_text = "\n".join([f"{m['sender']}: {m['text']}" for m in past_messages[:-1]])
        rewrite_prompt = f"""Given the following chat history and the user's latest question, rewrite the latest question into a standalone query that includes all necessary context from the history. If it is already standalone, return it as is. Do not answer the question, just rewrite it.
History:
{history_text}
Latest Question: {question}
Standalone Query:"""
        try:
            gemini_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-3.5-flash:generateContent?key={GEMINI_API_KEY}"
            payload = {"contents": [{"parts": [{"text": rewrite_prompt}]}]}
            async with httpx.AsyncClient(timeout=10.0) as client:
                res = await client.post(gemini_url, json=payload, headers={"Content-Type": "application/json"})
                if res.status_code == 200:
                    data = res.json()
                    if "candidates" in data and data["candidates"]:
                        rewritten = data["candidates"][0]["content"]["parts"][0]["text"].strip()
                        if rewritten:
                            search_query = rewritten
        except Exception as e:
            print("Query rewrite failed:", e)

    # ChromaDB query with per-file filtering
    where_clause = {"device_id": user_id}
    if file_id:
        where_clause = {"$and": [{"device_id": user_id}, {"file_id": file_id}]}

    results = collection.query(
        query_texts=[search_query],
        n_results=10,
        where=where_clause
    )

    # Similarity Score Filtering & Context building with Citation
    relevant_chunks = []
    if results["documents"] and results["documents"][0]:
        for idx, doc in enumerate(results["documents"][0]):
            distance = results["distances"][0][idx]
            metadata = results["metadatas"][0][idx]
            filename = metadata.get("filename", "Unknown File")
            page_num = metadata.get("page_num", 1)
            
            # Distance threshold (lower is better for L2). Usually < 1.3 is good.
            if distance < 1.3:
                relevant_chunks.append(f"[Source: {filename} - Page {page_num}]\n{doc}")

    context = "\n\n---\n\n".join(relevant_chunks)

    # Prompt banao
    prompt = f"""You are a helpful document assistant. The "Context" provided below is your primary source of knowledge.

STRICT RULES:
1. If the user is just saying hello, greeting you, or asking who/what you are, introduce yourself politely as an AI Document Assistant.
2. For all other factual or analytical questions, answer ONLY based on what is written in the "Context" below. Do NOT use outside knowledge.
3. Every piece of context starts with [Source: filename]. If you use information from the context, you MUST append the source filename at the end of your answer, formatted exactly as: "Source: filename".
4. If the question requires information from the documents but the Context is empty or doesn't contain the answer, clearly state: "This information was not found in the uploaded files. Please ensure you have uploaded the relevant document."

Context:
{context}

Question: {question}

Answer:"""

    # Streaming generator
    async def generate():
        # Check system state (files processing, empty, etc.)
        user_files = list(files_collection.find({"device_id": user_id}))
        ready = [f for f in user_files if f.get("status") == "ready"]
        
        system_error_answer = None
        if user_files and not ready:
            processing = [f for f in user_files if f.get("status") == "processing"]
            errored = [f for f in user_files if f.get("status") == "error"]
            
            if processing:
                pct = processing[0].get("progress", 0)
                system_error_answer = f"⏳ Your file is still being processed ({pct}% done). Please wait a moment — you can ask questions once it's ready!"
            elif errored:
                err_msg = errored[0].get("error_message", "Unknown error")
                system_error_answer = f"❌ File processing failed: {err_msg}\n\nPlease delete the file and try uploading it again."
                
        if system_error_answer:
            messages_collection.insert_one({
                "session_id": session_id,
                "sender": "ai",
                "text": system_error_answer,
                "timestamp": datetime.utcnow()
            })
            yield f"data: {json.dumps({'token': system_error_answer})}\n\n"
            yield f"data: {json.dumps({'done': True})}\n\n"
            return

        full_answer = ""
        try:
            gemini_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-3.5-flash:streamGenerateContent?alt=sse&key={GEMINI_API_KEY}"
            
            payload = {
                "contents": [{
                    "parts": [{"text": prompt}]
                }]
            }

            async with httpx.AsyncClient(timeout=60.0) as client:
                async with client.stream("POST", gemini_url, json=payload, headers={"Content-Type": "application/json"}) as response:
                    if response.status_code != 200:
                        await response.aread()
                        error_text = response.text
                        raise Exception(f"Gemini API Error ({response.status_code}): {error_text}")

                    async for line in response.aiter_lines():
                        if line.startswith("data: "):
                            data_str = line[6:]
                            if data_str.strip() == "[DONE]":
                                break
                            try:
                                data = json.loads(data_str)
                                if "candidates" in data and data["candidates"]:
                                    parts = data["candidates"][0].get("content", {}).get("parts", [])
                                    if parts and "text" in parts[0]:
                                        token = parts[0]["text"]
                                        full_answer += token
                                        yield f"data: {json.dumps({'token': token})}\n\n"
                                    else:
                                        # Handle safety block or empty parts
                                        finish_reason = data["candidates"][0].get("finishReason", "")
                                        if finish_reason:
                                            blocked_msg = f"\n\n[Response blocked by AI safety filters: {finish_reason}]"
                                            full_answer += blocked_msg
                                            yield f"data: {json.dumps({'token': blocked_msg})}\n\n"
                            except:
                                continue

            if not full_answer.strip():
                full_answer = "Sorry, the AI returned an empty response. It might have been blocked by safety filters or an internal error occurred."
                yield f"data: {json.dumps({'token': full_answer})}\n\n"

            # Final answer MongoDB mein save karo
            messages_collection.insert_one({
                "session_id": session_id,
                "sender": "ai",
                "text": full_answer,
                "timestamp": datetime.utcnow()
            })

            yield f"data: {json.dumps({'done': True})}\n\n"

        except Exception as e:
            error_str = str(e)
            if "429" in error_str or "RESOURCE_EXHAUSTED" in error_str or "quota" in error_str.lower():
                error_msg = "⏳ The AI service has hit its daily free-tier limit. Please try again in a few minutes, or later today once the quota resets."
            else:
                error_msg = "⚠️ An internal error occurred while processing your request. Please try again."
            print(f"System Error in chat: {e}")
            messages_collection.insert_one({
                "session_id": session_id,
                "sender": "ai",
                "text": error_msg,
                "timestamp": datetime.utcnow()
            })
            yield f"data: {json.dumps({'token': error_msg, 'done': True})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")